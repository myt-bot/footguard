"""Build the competition implementation edition of the FootGuard design plan."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "足安智垫_双足版本详细方案_竞赛实现版.docx"

BLUE = "167D78"
DARK = "163331"
LIGHT = "E8F4F2"
GRAY = "F2F4F7"
RED = "A33A3A"


def set_font(run, name: str = "Noto Sans CJK SC", size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_table_header(table) -> None:
    for cell in table.rows[0].cells:
        set_cell_fill(cell, LIGHT)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_fixed(table, widths)
    shade_table_header(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_font(run, size=9)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(doc: Document, title: str, body: str, danger: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, "FDECEC" if danger else LIGHT)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}  ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(RED if danger else BLUE)
    paragraph.add_run(body)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 28, DARK, 0, 12),
        ("Subtitle", 13, BLUE, 0, 16),
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "FootGuard 足安智垫｜竞赛实现版"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        set_font(run, size=8.5)
        run.font.color.rgb = RGBColor(110, 125, 123)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("工程原型｜非医疗诊断")
    set_font(run, size=8.5)
    run.font.color.rgb = RGBColor(110, 125, 123)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("“足安智垫”")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("基于双足多源感知、个人基线、云端大模型与闭环干预的\n糖尿病足风险辅助监测系统")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("竞赛实现版 V4.0｜2026 年 7 月 26 日\n布局：双足 6P4T｜主控：ESP32-S3｜App：Flutter｜后端：FastAPI")

    add_callout(
        doc,
        "当前结论",
        "核心端到端闭环已经形成。本版本只把代码、实物或测试可证明的能力写为“已实现”；"
        "固定前的阶段性实测可证明通道响应和主要风险方向，但因位置、接触及手工左右结构存在"
        "不确定性，不用于锁定最终标定值。胶固化后的复测仅用于复核传感器映射、基线重复性和工程阈值，"
        "不用推翻通信、历史、基线入口、"
        "云端解释、马达 ACK 与效果评价结构。",
    )

    doc.add_heading("1. 项目定位与边界", level=1)
    doc.add_paragraph(
        "足安智垫面向糖尿病足日常管理场景，通过左右脚关键区域压力与温度趋势、双足差异和持续时间，"
        "给出风险辅助提示。项目的差异化不是追求几十点连续压力成像，而是以低成本关键点感知完成"
        "“发现—解释—提醒—复测—评价”的闭环。"
    )
    add_bullets(
        doc,
        [
            "目标用户：需要关注足底异常受力、左右温差和日常足部状态的人群；当前仅作竞赛工程验证。",
            "产品边界：不诊断糖尿病足、不替代医生、不把竞赛阈值宣传为临床标准。",
            "医学相关性：IWGDF 2023 建议可考虑指导中高风险人群监测足部皮肤温度；本原型据此关注左右同区温差，但采用更短时间只为安全演示，不能等同临床判定。",
            "赛题匹配：使用 ESP32-S3 完成传感器采集、预处理、BLE 和执行控制，App 与云端大模型构成完整 AIoT 服务链。",
        ],
    )

    doc.add_heading("2. 系统总体架构", level=1)
    add_table(
        doc,
        ["层级", "当前实现", "职责"],
        [
            ["左/右鞋端", "2×ESP32-S3、每脚 6 FSR + 4 NTC + MPU6050 + 马达", "5 Hz 采集、质量标记、BLE、命令执行与 ACK"],
            ["移动端", "Flutter Android App", "双 BLE、TimeSync、配对、实时展示、上传、历史与设置"],
            ["规则后端", "FastAPI + SQLite", "个人基线、持续风险、命令、恢复评价和历史"],
            ["云端解释", "DeepSeek OpenAI-compatible API", "把结构化风险转成非诊断性解释；失败时模板降级"],
        ],
        [1600, 3000, 4760],
    )
    doc.add_paragraph(
        "主链路：双足传感器 → ESP32-S3 → BLE → App → FastAPI → 确定性风险 → DeepSeek/模板解释 → "
        "白名单马达命令 → App → 目标脚 ESP32-S3 → AckEvent → 后端 → 干预后数据 → 效果评价。"
    )

    doc.add_heading("3. 硬件与点位", level=1)
    add_table(
        doc,
        ["通道", "区域", "用途"],
        [
            ["P1", "拇趾/最前端", "拇趾与最前端负重趋势"],
            ["P2", "前掌外侧", "前掌外侧相对负重"],
            ["P3", "前掌中央", "前掌中央相对负重"],
            ["P4", "前掌内侧", "第一跖骨附近相对负重"],
            ["P5", "中足中央", "中足支撑趋势"],
            ["P6", "足跟中央", "足跟支撑趋势"],
            ["T1～T4", "前掌外侧、拇趾附近、足跟、中足", "左右同区温差与个人趋势"],
        ],
        [1400, 3000, 4960],
    )
    add_callout(
        doc,
        "电气说明",
        "FSR 通过 47 kΩ 分压读入 ADC，输出 0～4095 原始值后归一化；NTC 使用 10K B3950 分压和 Beta 公式换算温度。"
        "马达现有原型由 GPIO13 控制；正式可穿戴版本应使用驱动器件和反向保护，避免直接由 GPIO 承担马达电流。",
        danger=True,
    )

    doc.add_heading("4. 固件、BLE 与安全执行", level=1)
    add_bullets(
        doc,
        [
            "左右脚分别构建，设备身份为 foot_left_001 / foot_right_001，广播名为 FootGuard-L / FootGuard-R。",
            "固定 60 字节 SensorData 包携带布局版本、序号、同步时间、6 路压力、4 路温度、IMU、电量占位和质量位。",
            "DeviceCommand 经过 Schema、目标、有效期和 command_id 校验后进入异步执行队列。",
            "固件缓存 AckEvent；断线重连或重复写入时可重放 ACK，但不会再次执行同一命令。",
        ],
    )

    doc.add_heading("5. App 与双足同步", level=1)
    add_bullets(
        doc,
        [
            "双设备扫描、连接、订阅、自动重连和设备身份校验。",
            "TimeSync 后按序号与时间进行左右帧配对，短时保持最近完整配对，避免页面因到达抖动闪烁。",
            "无效传感器显示为“--”，不会用 0℃或漂移值制造热区。",
            "实时页展示双足压力热力图、温差、负重偏向、风险、DeepSeek 解释和马达 ACK。",
            "设置页展示基线学习进度和重新校准入口；历史页展示风险、震动和干预前后效果。",
        ],
    )

    doc.add_heading("6. 个人基线与风险算法", level=1)
    doc.add_heading("6.1 基线", level=2)
    doc.add_paragraph(
        "体验者自然站立、双脚平行且与肩同宽。后端从稳定承重且温差相对稳定的配对帧中选取候选，"
        "达到默认 15 个样本后以稳健中位数锁定个人压力分布和正常左右差异。换人、移动传感器或改变电路后必须重新校准。"
    )
    doc.add_heading("6.2 当前确定性风险", level=2)
    add_table(
        doc,
        ["风险", "判定基础", "验证状态"],
        [
            ["left/right_load_bias", "双足总载荷归一化差异 + 持续时间", "已实现，固定后复测阈值"],
            ["forefoot_high", "P1～P4 前掌占比相对个人基线增量", "已实现，固定后复测灵敏度"],
            ["temperature_asymmetry", "左右同区温差扣除个人基线", "已实现，阈值仅作工程演示"],
            ["data_incomplete", "单侧、同步或传感器质量异常", "已实现，不下发马达"],
            ["步态/冲击增强", "MPU + 压力周期与冲击特征", "数据链已接入，尚待可靠行走数据验证"],
        ],
        [2200, 4660, 2500],
    )
    add_callout(
        doc,
        "关于“复杂算法”",
        "当前数据量不足以证明 XGBoost、TCN 或 AutoEncoder 的泛化效果。今天不为了显得复杂而加入未经验证的模型；"
        "固定后先保证基线和规则可重复，后续再按参与者分组采集数据并报告准确率、召回率、F1 与消融结果。",
    )

    doc.add_heading("7. 分级震动与效果评价", level=1)
    add_table(
        doc,
        ["等级", "状态", "动作"],
        [
            ["0", "正常", "不震动"],
            ["1", "关注，约持续 3 秒", "只显示提示，不震动"],
            ["2", "警告，约持续 6 秒", "风险侧 double，双短震 800 ms"],
            ["3", "持续异常，约持续 10 秒", "风险侧 long，长震 1500 ms"],
        ],
        [1200, 3760, 4400],
    )
    doc.add_paragraph(
        "同一事件可从等级 2 升级至等级 3，各等级只执行一次。事件恢复后，系统优先关联真实 executed ACK，"
        "比较干预前后 load_diff 等指标，输出 effective、partial、ineffective 或无法评价，并记录恢复时间。"
    )

    doc.add_heading("8. 云端大模型", level=1)
    add_bullets(
        doc,
        [
            "输入为风险类型、侧别、等级、持续时间、负重差和区域分析，不直接发送连续原始流。",
            "输出采用固定 Schema：risk_level、explanation、advice、target、candidate_pattern。",
            "服务端强制以本地规则覆盖 target 和 candidate_pattern；大模型不能自由驱动硬件。",
            "DeepSeek 超时、401、格式错误或无密钥时自动使用 mock-risk-advisor 模板。",
        ],
    )

    doc.add_heading("9. 历史页和干预前后展示", level=1)
    add_bullets(
        doc,
        [
            "按时间列出风险类型、侧别、等级、开始/结束、持续时间和是否仍进行中。",
            "展示震动目标侧、模式、设备执行状态和错误码。",
            "展开事件可查看干预前后 load_diff、改善百分比、恢复时间和效果结论。",
            "缺少 ACK 或恢复窗口数据时明确显示无法评价，不生成虚假改善。",
        ],
    )

    doc.add_heading("10. 当前完成度与待验证项", level=1)
    add_table(
        doc,
        ["模块", "状态", "提交前动作"],
        [
            ["双足传感与 BLE", "已实现", "胶固化后检查映射、空载和连续连接"],
            ["个人基线与重校准", "代码已实现", "App/后端全量测试 + 真机一次"],
            ["规则与分级马达", "代码已实现", "验证 double/long、ACK 和去重"],
            ["DeepSeek 解释", "已实现", "确认密钥、降级模板和中文编码"],
            ["历史与效果", "已实现", "产生一次真实风险—恢复事件"],
            ["MPU 步态融合", "初步实现", "不作为初赛核心已验证能力"],
            ["真实电量", "未实现", "初赛可不加；页面不可宣称 95% 为真实电量"],
            ["临床有效性", "未验证", "明确工程原型边界"],
        ],
        [2200, 2300, 4860],
    )

    doc.add_heading("11. 胶固化后的最终复测顺序", level=1)
    add_numbered(
        doc,
        [
            "双脚离地，检查六路压力空载与四路温度稳定性；逐点验证 P1～P6、T1～T4 映射。",
            "重新校准个人基线，连续完成三次相同站姿，比较结果能否重复。",
            "验证正常站立不误报，再验证偏左、偏右、前掌持续高载和恢复。",
            "若结果偏差大，先排查接触与固定，再只调整 config.py 中工程阈值；每次调整后复测相同场景。",
            "连接马达，验证等级 2 双短震、等级 3 长震、目标侧、ACK、去重和过期。",
            "打开历史页，确认最新事件有干预前后值、恢复时间和效果结论。",
            "最后录制温度趋势、DeepSeek 解释和离线降级，不再临时加入未经验证的模型。",
        ],
    )

    doc.add_heading("12. 演示口径", level=1)
    doc.add_paragraph(
        "演示顺序建议为：双设备连接 → 重新校准 → 正常状态 → 持续偏载/前掌高载 → "
        "等级 2 双短震 → 可选等级 3 长震 → 调整重心恢复 → 历史页干预前后对比 → "
        "手掌安全捂热单侧 NTC → DeepSeek 解释与模板降级。"
    )
    add_callout(
        doc,
        "核心答辩句",
        "我们没有让大模型直接控制医疗相关动作。ESP32-S3 完成真实感知、BLE 和执行，"
        "确定性规则负责风险与白名单命令，大模型只负责通俗解释；提醒后用传感器继续验证是否改善。",
    )

    doc.add_heading("13. 竞赛要求对应", level=1)
    add_table(
        doc,
        ["要求", "对应实现"],
        [
            ["乐鑫指定芯片", "左右两侧均以 ESP32-S3 为主控"],
            ["多维感知", "压力、温度、IMU 与设备质量状态"],
            ["连接与服务", "BLE 双设备 + Flutter 移动网关 + FastAPI"],
            ["云端大模型", "DeepSeek 结构化解释与本地降级"],
            ["交互与执行", "App、目标侧马达、ACK 与历史闭环"],
            ["完整性与可行性", "真实硬件、协议、测试、Mock 备用和明确边界"],
        ],
        [2600, 6760],
    )

    doc.add_heading("14. 参考资料", level=1)
    refs = [
        "乐鑫科技：2026 年全国大学生物联网设计竞赛乐鑫命题，https://www.espressif.com/en/ecosystem/education/competition/iot",
        "IWGDF Guidelines (2023 update)，https://iwgdfguidelines.org/guidelines-2023/",
        "Bus SA 等：Guidelines on the prevention of foot ulcers in persons with diabetes (IWGDF 2023 update)，Diabetes/Metabolism Research and Reviews，2024。",
    ]
    for index, reference in enumerate(refs, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(f"{index}.  {reference}")
    doc.add_paragraph(
        "说明：IWGDF 的对应区域温差建议针对每日监测并要求连续两天；本项目当前 2.0℃、秒级持续判断仅用于竞赛工程演示，"
        "不可写成临床诊断阈值或预防效果证明。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附录 A：提交前验收清单", level=1)
    add_bullets(
        doc,
        [
            "后端测试全部通过；Flutter analyze 和 test 全部通过；左右固件均可构建。",
            "App 能显示基线状态并完成确认式重置。",
            "等级 2/3 的模式、时长、目标侧和 ACK 与后端记录一致。",
            "历史页显示至少一个真实事件的前后效果。",
            "DeepSeek 在线与无密钥降级各验证一次。",
            "README、详细方案、演示视频口径与真实功能一致。",
            "演示中不展示 API Key、个人隐私或无法解释的调试日志。",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
